from docx import Document as DocxDocument

from ..schemas.document import Document
from ..utils.text_splitter import TextSplitter



class DocxSplitter:
    def __init__(self, docx_path: str, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.docx_path = docx_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def __get_content(self) -> str:
        try:
            document = DocxDocument(self.docx_path)
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error loading .docx file: {e}")
            return ""

    def split_documents(self) -> list[Document]:
        content = self.__get_content()
        splitter = TextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
        chunks = splitter.split(content)
        documents = [Document(content=chunk, metadata={'source': self.docx_path}) for chunk in chunks]
        return documents